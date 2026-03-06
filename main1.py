import os
import re
import spacy
from docx import Document
from datetime import datetime

INPUT_FOLDER = "input"
OUTPUT_FOLDER = "output"
LOG_FILE = "anonymization_log.txt"

nlp = spacy.load("sv_core_news_sm")

PERSONNUMMER_REGEX = r"\b(19|20)?\d{6}[- ]?\d{4}\b"
EMAIL_REGEX = r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"
PHONE_REGEX = r"\b(\+46\s?|0)[0-9\s\-]{8,12}\b"

# fångar:
# Marcus Lampinen
# M Lampinen
# E-L Westberg
NAME_REGEX = r"\b([A-ZÅÄÖ]\.?|[A-ZÅÄÖ][a-zåäö]+|[A-ZÅÄÖ]-[A-ZÅÄÖ])\s[A-ZÅÄÖ][a-zåäö\-]+\b"

# fångar kvarvarande ensamma namn
"SINGLE_NAME_REGEX = r"\b[A-ZÅÄÖ][a-zåäö]{2,}\b""


def log(message):

    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"{datetime.now()} - {message}\n")


def anonymize_text(text):

    original = text

    text = re.sub(PERSONNUMMER_REGEX, "[PERSONNUMMER]", text)
    text = re.sub(EMAIL_REGEX, "[EMAIL]", text)
    text = re.sub(PHONE_REGEX, "[TELEFON]", text)

    doc = nlp(text)

    for ent in doc.ents:
        if ent.label_ == "PER":
            text = text.replace(ent.text, "[PERSON]")

    text = re.sub(NAME_REGEX, "[PERSON]", text)

    # anonymisera ensamma namn
    "text = re.sub(SINGLE_NAME_REGEX, "[PERSON]", text)"

    if text != original:
        log("Text anonymiserad")

    return text


def detect_persons(text):

    persons = set()

    doc = nlp(text)

    for ent in doc.ents:
        if ent.label_ == "PER":
            persons.add(ent.text)

    regex_names = re.findall(NAME_REGEX, text)

    for name in regex_names:
        persons.add(name)

    return persons


def anonymize_paragraph(paragraph):

    text = paragraph.text

    if not text.strip():
        return

    anonymized = anonymize_text(text)

    if anonymized != text:

        for run in paragraph.runs[::-1]:
            paragraph._element.remove(run._element)

        paragraph.add_run(anonymized)


def process_paragraphs(doc):

    for paragraph in doc.paragraphs:
        anonymize_paragraph(paragraph)


def process_tables(doc):

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:
                    anonymize_paragraph(paragraph)


def process_headers_footers(doc):

    for section in doc.sections:

        header = section.header
        footer = section.footer

        for paragraph in header.paragraphs:
            anonymize_paragraph(paragraph)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        anonymize_paragraph(paragraph)

        for paragraph in footer.paragraphs:
            anonymize_paragraph(paragraph)

        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        anonymize_paragraph(paragraph)


def process_shapes(doc):

    try:

        for shape in doc.inline_shapes:

            if hasattr(shape, "text_frame"):

                for paragraph in shape.text_frame.paragraphs:
                    anonymize_paragraph(paragraph)

        log("Shapes analyserade")

    except:
        log("Shapes kunde inte analyseras")


def clean_metadata(doc):

    props = doc.core_properties

    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.title = ""
    props.subject = ""

    log("Metadata rensad")


def remove_comments(doc):

    try:

        comments_part = doc.part._comments_part

        if comments_part is not None:
            comments_part._element.clear()

            log("Word kommentarer borttagna")

    except:
        log("Inga kommentarer hittades")


def remove_track_changes(doc):

    try:

        body = doc.part._element.body

        for element in body.xpath(".//w:ins | .//w:del"):
            element.getparent().remove(element)

        log("Track changes rensade")

    except:
        log("Track changes kunde inte analyseras")


def scan_document_for_persons(doc):

    persons = set()

    for paragraph in doc.paragraphs:
        persons.update(detect_persons(paragraph.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    persons.update(detect_persons(paragraph.text))

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            persons.update(detect_persons(paragraph.text))

        for paragraph in section.footer.paragraphs:
            persons.update(detect_persons(paragraph.text))

    return persons


def clean_filename(filename):

    name = filename.lower()

    name = re.sub(r"[åäö]", "", name)
    name = re.sub(r"[^a-z0-9_.]", "_", name)

    return "anon_" + name


def anonymize_docx(input_file, output_file):

    doc = Document(input_file)

    clean_metadata(doc)

    remove_comments(doc)

    remove_track_changes(doc)

    process_paragraphs(doc)

    process_tables(doc)

    process_headers_footers(doc)

    process_shapes(doc)

    doc.save(output_file)


def process_documents():

    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)

    files = os.listdir(INPUT_FOLDER)

    all_persons = set()

    for file in files:

        if file.endswith(".docx"):

            input_path = os.path.join(INPUT_FOLDER, file)

            doc = Document(input_path)

            persons = scan_document_for_persons(doc)

            all_persons.update(persons)

    if all_persons:

        print("\nIdentifierade personer:\n")

        for p in sorted(all_persons):
            print(p)

        print("\nStartar anonymisering...\n")

    for file in files:

        if file.endswith(".docx"):

            input_path = os.path.join(INPUT_FOLDER, file)

            output_name = clean_filename(file)

            output_path = os.path.join(OUTPUT_FOLDER, output_name)

            print(f"Bearbetar {file}")

            log(f"Start anonymisering {file}")

            anonymize_docx(input_path, output_path)

            log(f"Klar {file}")

            print(f"Klar {file}")


if __name__ == "__main__":

    print("Startar anonymisering")

    process_documents()

    print("\nAlla dokument färdiga\n")