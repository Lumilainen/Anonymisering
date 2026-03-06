import os
import re
import spacy
from docx import Document
from datetime import datetime

INPUT_FOLDER = "input"
OUTPUT_FOLDER = "output"
LOG_FILE = "anonymization_log.txt"

nlp = spacy.load("sv_core_news_sm")

PERSONNUMMER_REGEX = r"\b(19|20)?\d{6}[- ]?\d{4}\b"
EMAIL_REGEX = r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"
PHONE_REGEX = r"\b(\+46\s?|0)[0-9\s\-]{8,12}\b"

# Förnamn Efternamn
NAME_REGEX = r"\b([A-ZÅÄÖ][a-zåäö]+(?:-[A-ZÅÄÖ][a-zåäö]+)?)\s([A-ZÅÄÖ][a-zåäö]+)\b"

# Initial Efternamn
INITIAL_NAME_REGEX = r"\b([A-ZÅÄÖ](?:-[A-ZÅÄÖ])?)\s([A-ZÅÄÖ][a-zåäö]+)\b"

EXCLUDED_WORDS = {
    "Informationssäkerhetschef",
    "Rapportering",
    "Ansvarsområde",
    "Befattning",
    "Datum",
    "Dokument"
}


def log(message):

    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"{datetime.now()} - {message}\n")


def normalize_text(text):

    text = re.sub(r"\s+", " ", text)

    return text.strip()


def detect_persons(text):

    persons = set()

    text = normalize_text(text)

    doc = nlp(text)

    # spaCy persondetektion
    for ent in doc.ents:

        if ent.label_ == "PER":

            name = ent.text.strip()

            if name not in EXCLUDED_WORDS:

                persons.add(name)

                parts = name.split()

                if len(parts) > 1:
                    persons.add(parts[-1])  # lägg till efternamn

    # regex för förnamn efternamn
    matches = re.findall(NAME_REGEX, text)

    for match in matches:

        full = " ".join(match)

        persons.add(full)

        persons.add(match[1])  # efternamn

    # regex för initialnamn
    matches = re.findall(INITIAL_NAME_REGEX, text)

    for match in matches:

        full = " ".join(match)

        persons.add(full)

        persons.add(match[1])

    return persons


def anonymize_text(text, persons):

    original = text

    text = re.sub(PERSONNUMMER_REGEX, "[PERSONNUMMER]", text)
    text = re.sub(EMAIL_REGEX, "[EMAIL]", text)
    text = re.sub(PHONE_REGEX, "[TELEFON]", text)

    for name in sorted(persons, key=len, reverse=True):

        escaped = re.escape(name)

        text = re.sub(rf"\b{escaped}\b", "[PERSON]", text)

    if text != original:
        log("Text anonymiserad")

    return text


def anonymize_paragraph(paragraph, persons):

    text = paragraph.text

    if not text.strip():
        return

    anonymized = anonymize_text(text, persons)

    if anonymized != text:

        for run in paragraph.runs[::-1]:
            paragraph._element.remove(run._element)

        paragraph.add_run(anonymized)


def process_paragraphs(doc, persons):

    for paragraph in doc.paragraphs:

        anonymize_paragraph(paragraph, persons)


def process_tables(doc, persons):

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    anonymize_paragraph(paragraph, persons)


def process_headers_footers(doc, persons):

    for section in doc.sections:

        header = section.header
        footer = section.footer

        for paragraph in header.paragraphs:

            anonymize_paragraph(paragraph, persons)

        for table in header.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        anonymize_paragraph(paragraph, persons)

        for paragraph in footer.paragraphs:

            anonymize_paragraph(paragraph, persons)

        for table in footer.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        anonymize_paragraph(paragraph, persons)


def process_shapes(doc, persons):

    try:

        for shape in doc.inline_shapes:

            if hasattr(shape, "text_frame"):

                for paragraph in shape.text_frame.paragraphs:

                    anonymize_paragraph(paragraph, persons)

    except:

        log("Shapes kunde inte analyseras")


def clean_metadata(doc):

    props = doc.core_properties

    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.title = ""
    props.subject = ""

    log("Metadata rensad")


def remove_comments(doc):

    try:

        comments_part = doc.part._comments_part

        if comments_part is not None:

            comments_part._element.clear()

            log("Word kommentarer borttagna")

    except:

        log("Inga kommentarer hittades")


def remove_track_changes(doc):

    try:

        body = doc.part._element.body

        for element in body.xpath(".//w:ins | .//w:del"):

            element.getparent().remove(element)

        log("Track changes rensade")

    except:

        log("Track changes kunde inte analyseras")


def scan_document_for_persons(doc):

    persons = set()

    for paragraph in doc.paragraphs:

        persons.update(detect_persons(paragraph.text))

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    persons.update(detect_persons(paragraph.text))

    for section in doc.sections:

        header = section.header
        footer = section.footer

        for paragraph in header.paragraphs:

            persons.update(detect_persons(paragraph.text))

        for table in header.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        persons.update(detect_persons(paragraph.text))

        for paragraph in footer.paragraphs:

            persons.update(detect_persons(paragraph.text))

        for table in footer.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        persons.update(detect_persons(paragraph.text))

    return persons


def clean_filename(filename):

    name = filename.lower()

    name = re.sub(r"[åäö]", "", name)

    name = re.sub(r"[^a-z0-9_.]", "_", name)

    return "anon_" + name


def anonymize_docx(input_file, output_file, persons):

    doc = Document(input_file)

    clean_metadata(doc)

    remove_comments(doc)

    remove_track_changes(doc)

    process_paragraphs(doc, persons)

    process_tables(doc, persons)

    process_headers_footers(doc, persons)

    process_shapes(doc, persons)

    doc.save(output_file)


def process_documents():

    if not os.path.exists(OUTPUT_FOLDER):

        os.makedirs(OUTPUT_FOLDER)

    files = os.listdir(INPUT_FOLDER)

    all_persons = set()

    for file in files:

        if file.endswith(".docx"):

            input_path = os.path.join(INPUT_FOLDER, file)

            doc = Document(input_path)

            persons = scan_document_for_persons(doc)

            all_persons.update(persons)

    if all_persons:

        print("\nIdentifierade personer:\n")

        for p in sorted(all_persons):

            print(p)

        print("\nStartar anonymisering...\n")

    for file in files:

        if file.endswith(".docx"):

            input_path = os.path.join(INPUT_FOLDER, file)

            output_name = clean_filename(file)

            output_path = os.path.join(OUTPUT_FOLDER, output_name)

            print(f"Bearbetar {file}")

            anonymize_docx(input_path, output_path, all_persons)

            print(f"Klar {file}")


if __name__ == "__main__":

    print("Startar anonymisering")

    process_documents()

    print("\nAlla dokument färdiga\n")